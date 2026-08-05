# -*- coding: utf-8 -*-
"""Convierte el informe HTML a un documento Word editable (.docx)."""
import re, sys
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = Path(__file__).parent
AZUL = RGBColor(0x0B, 0x3C, 0x5D)
AZUL_HEX = "0B3C5D"
GRIS = RGBColor(0x4A, 0x5A, 0x63)
ANCHO_UTIL_CM = 17.2


def sombrear(celda, color_hex):
    propiedades = celda._tc.get_or_add_tcPr()
    sombra = OxmlElement("w:shd")
    sombra.set(qn("w:val"), "clear")
    sombra.set(qn("w:fill"), color_hex)
    propiedades.append(sombra)


def sin_bordes(tabla):
    propiedades = tabla._tbl.tblPr
    bordes = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elemento = OxmlElement(f"w:{lado}")
        elemento.set(qn("w:val"), "none")
        bordes.append(elemento)
    propiedades.append(bordes)


def barra_izquierda(tabla, color_hex):
    propiedades = tabla._tbl.tblPr
    bordes = OxmlElement("w:tblBorders")
    for lado in ("top", "bottom", "right", "insideH", "insideV"):
        elemento = OxmlElement(f"w:{lado}")
        elemento.set(qn("w:val"), "none")
        bordes.append(elemento)
    izquierda = OxmlElement("w:left")
    izquierda.set(qn("w:val"), "single")
    izquierda.set(qn("w:sz"), "18")
    izquierda.set(qn("w:color"), color_hex)
    bordes.append(izquierda)
    propiedades.append(bordes)


def escribir_inline(parrafo, elemento, negrita=False, cursiva=False, tamano=None, color=None):
    """Vuelca el contenido de un nodo HTML en un párrafo de Word conservando b / i / br."""
    for hijo in elemento.children:
        if isinstance(hijo, NavigableString):
            texto = re.sub(r"\s+", " ", str(hijo))
            if not texto.strip() and texto != " ":
                continue
            corrida = parrafo.add_run(texto.replace(" ", " "))
            corrida.bold, corrida.italic = negrita, cursiva
            if tamano:
                corrida.font.size = tamano
            if color:
                corrida.font.color.rgb = color
        elif hijo.name == "br":
            parrafo.add_run().add_break()
        elif hijo.name in ("b", "strong"):
            escribir_inline(parrafo, hijo, True, cursiva, tamano, color)
        elif hijo.name in ("i", "em"):
            escribir_inline(parrafo, hijo, negrita, True, tamano, color)
        else:
            escribir_inline(parrafo, hijo, negrita, cursiva, tamano, color)


def agregar_tabla(documento, nodo):
    filas = nodo.find_all("tr")
    if not filas:
        return
    columnas = len(filas[0].find_all(["th", "td"]))
    es_kpi = "kpi" in (nodo.get("class") or [])

    tabla = documento.add_table(rows=0, cols=columnas)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    anchos = None
    grupo = nodo.find("colgroup")
    if grupo:
        porcentajes = [int(c.get("width", "0").rstrip("%")) for c in grupo.find_all("col")]
        if sum(porcentajes) > 0:
            anchos = [Cm(ANCHO_UTIL_CM * p / 100) for p in porcentajes]

    for indice_fila, fila_html in enumerate(filas):
        celdas_html = fila_html.find_all(["th", "td"])
        fila = tabla.add_row()
        for indice_columna, celda_html in enumerate(celdas_html):
            if indice_columna >= columnas:
                continue
            celda = fila.cells[indice_columna]
            celda.text = ""
            parrafo = celda.paragraphs[0]
            parrafo.paragraph_format.space_before = Pt(1.5)
            parrafo.paragraph_format.space_after = Pt(1.5)

            if es_kpi:
                parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                valor = celda_html.find("span", class_="valor")
                etiqueta = celda_html.find("span", class_="etiqueta")
                if valor is not None:
                    corrida = parrafo.add_run(valor.get_text(strip=True))
                    corrida.bold = True
                    corrida.font.size = Pt(14)
                    corrida.font.color.rgb = AZUL
                if etiqueta is not None:
                    parrafo.add_run().add_break()
                    corrida = parrafo.add_run(etiqueta.get_text(strip=True))
                    corrida.font.size = Pt(7.5)
                    corrida.font.color.rgb = GRIS
                sombrear(celda, "F7FAFB")
                continue

            es_encabezado = celda_html.name == "th"
            escribir_inline(parrafo, celda_html, negrita=es_encabezado, tamano=Pt(8.5),
                            color=RGBColor(0xFF, 0xFF, 0xFF) if es_encabezado else None)
            if es_encabezado:
                sombrear(celda, AZUL_HEX)
            elif indice_fila % 2 == 0:
                sombrear(celda, "F2F6F8")

            if anchos and indice_columna < len(anchos):
                celda.width = anchos[indice_columna]

    documento.add_paragraph().paragraph_format.space_after = Pt(2)


def agregar_recuadro(documento, nodo, color_barra, color_fondo):
    tabla = documento.add_table(rows=1, cols=1)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    celda = tabla.cell(0, 0)
    celda.text = ""
    parrafo = celda.paragraphs[0]
    parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    escribir_inline(parrafo, nodo, tamano=Pt(9))
    sombrear(celda, color_fondo)
    barra_izquierda(tabla, color_barra)
    documento.add_paragraph().paragraph_format.space_after = Pt(2)


def convertir(ruta_html, ruta_docx):
    sopa = BeautifulSoup(Path(ruta_html).read_text(encoding="utf-8"), "lxml")
    documento = Document()

    seccion = documento.sections[0]
    seccion.page_height, seccion.page_width = Cm(29.7), Cm(21)
    seccion.top_margin, seccion.bottom_margin = Cm(1.7), Cm(1.6)
    seccion.left_margin = seccion.right_margin = Cm(1.9)

    normal = documento.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1

    # Pie de página con numeración automática.
    pie = seccion.footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corrida = pie.add_run("Evaluación Final — Diplomado en Data Engineer · USACH · Valeria Luna · página ")
    corrida.font.size = Pt(7.5)
    corrida.font.color.rgb = GRIS
    campo = OxmlElement("w:fldSimple")
    campo.set(qn("w:instr"), "PAGE")
    pie._p.append(campo)

    cuerpo = sopa.body
    for nodo in cuerpo.find_all(recursive=False):
        clases = nodo.get("class") or []

        if nodo.get("id") == "pie":
            continue

        if nodo.name == "div" and "portada" in clases:
            for hijo in nodo.find_all(["p", "h1"], recursive=False):
                if hijo.name == "h1":
                    parrafo = documento.add_paragraph()
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    escribir_inline(parrafo, hijo, negrita=True, tamano=Pt(18), color=AZUL)
                else:
                    parrafo = documento.add_paragraph()
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tamano = Pt(11) if "subtitulo" in (hijo.get("class") or []) else Pt(9)
                    escribir_inline(parrafo, hijo, tamano=tamano, color=GRIS)
            continue

        if nodo.name == "div" and "resumen" in clases:
            agregar_recuadro(documento, nodo, AZUL_HEX, "EEF4F7")
            continue
        if nodo.name == "div" and "destacado" in clases:
            agregar_recuadro(documento, nodo, "D98324", "FFF6E5")
            continue

        if nodo.name == "div" and "figura" in clases:
            imagen = nodo.find("img")
            if imagen is not None:
                ruta = Path(imagen["src"])
                if not ruta.is_absolute():
                    ruta = BASE / ruta
                if ruta.exists():
                    coincidencia = re.search(r"([\d.]+)pt", imagen.get("style", ""))
                    puntos = float(coincidencia.group(1)) if coincidencia else 460
                    parrafo = documento.add_paragraph()
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    parrafo.add_run().add_picture(str(ruta), width=Pt(puntos))
            continue

        if nodo.name in ("h1", "h2", "h3"):
            parrafo = documento.add_paragraph()
            parrafo.paragraph_format.space_before = Pt(10 if nodo.name == "h2" else 8)
            parrafo.paragraph_format.space_after = Pt(3)
            tamano = {"h1": Pt(16), "h2": Pt(12), "h3": Pt(10.5)}[nodo.name]
            escribir_inline(parrafo, nodo, negrita=True, tamano=tamano, color=AZUL)
            if nodo.name == "h2":
                bordes = OxmlElement("w:pBdr")
                inferior = OxmlElement("w:bottom")
                inferior.set(qn("w:val"), "single")
                inferior.set(qn("w:sz"), "8")
                inferior.set(qn("w:color"), AZUL_HEX)
                bordes.append(inferior)
                parrafo._p.get_or_add_pPr().append(bordes)
            continue

        if nodo.name == "p":
            estilo = "Normal"
            parrafo = documento.add_paragraph(style=estilo)
            parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if "pie-figura" in clases:
                parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                escribir_inline(parrafo, nodo, cursiva=True, tamano=Pt(8), color=GRIS)
            elif "nota" in clases:
                escribir_inline(parrafo, nodo, tamano=Pt(8.5), color=GRIS)
            else:
                escribir_inline(parrafo, nodo)
            continue

        if nodo.name == "ul":
            for elemento in nodo.find_all("li", recursive=False):
                parrafo = documento.add_paragraph(style="List Bullet")
                parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                parrafo.paragraph_format.space_after = Pt(2)
                escribir_inline(parrafo, elemento)
            continue

        if nodo.name == "table":
            agregar_tabla(documento, nodo)
            continue

    documento.save(ruta_docx)
    print("Word generado:", ruta_docx)


if __name__ == "__main__":
    convertir(sys.argv[1], sys.argv[2])
